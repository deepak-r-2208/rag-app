"""Extract plain text from uploaded files of various formats.

Supported: .txt, .md, .csv, .json, .pdf (text or scanned), .docx, and
common image formats via OCR (.png, .jpg, .jpeg, .webp, .tiff).
"""

import io
import json

from fastapi import UploadFile

SUPPORTED_EXTENSIONS = {
    "txt", "md", "csv", "json", "pdf", "docx", "png", "jpg", "jpeg", "webp", "tiff",
}

# If a PDF page yields fewer than this many characters of text per page on
# average, we assume it's a scan and fall back to OCR instead.
_SCANNED_PDF_CHAR_THRESHOLD = 20


class UnsupportedFormatError(ValueError):
    pass


def extension_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


async def extract_text(file: UploadFile) -> str:
    ext = extension_of(file.filename or "")
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(f"Unsupported file type: .{ext}")

    raw = await file.read()

    if ext in ("txt", "md"):
        return raw.decode("utf-8", errors="replace")

    if ext == "csv":
        return raw.decode("utf-8", errors="replace")

    if ext == "json":
        try:
            parsed = json.loads(raw.decode("utf-8", errors="replace"))
            return json.dumps(parsed, indent=2)
        except json.JSONDecodeError:
            return raw.decode("utf-8", errors="replace")

    if ext == "pdf":
        return _extract_pdf(raw)

    if ext == "docx":
        return _extract_docx(raw)

    # image formats -> OCR
    return _ocr_image_bytes(raw)


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    pages_text: list[str] = []
    needs_ocr_pages: list[int] = []

    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        pages_text.append(text)
        if len(text) < _SCANNED_PDF_CHAR_THRESHOLD:
            needs_ocr_pages.append(i)

    # If most pages look like scans, OCR the whole document (pdf2image + tesseract).
    if needs_ocr_pages and len(needs_ocr_pages) >= max(1, len(pages_text) // 2):
        try:
            ocr_text = _ocr_pdf(raw)
            if ocr_text.strip():
                return ocr_text
        except Exception:
            # Fall through to whatever text extraction produced, even if thin.
            pass

    return "\n\n".join(pages_text)


def _ocr_pdf(raw: bytes) -> str:
    import pytesseract
    from pdf2image import convert_from_bytes

    images = convert_from_bytes(raw)
    return "\n\n".join(pytesseract.image_to_string(img) for img in images)


def _extract_docx(raw: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def _ocr_image_bytes(raw: bytes) -> str:
    import pytesseract
    from PIL import Image

    image = Image.open(io.BytesIO(raw))
    return pytesseract.image_to_string(image)
